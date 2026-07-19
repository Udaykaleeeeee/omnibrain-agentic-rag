"""
DOCX document parser using python-docx.
Extracts text and embedded images from Word documents.

PROVENANCE NOTE for DOCX:
Since DOCX files don't have a native page concept, we create pseudo-pages:
1. Each major section (based on paragraph breaks and headings) becomes a "page"
2. Images are assigned to the page of their nearest preceding paragraph
3. Page numbers are assigned sequentially (1-indexed)
4. This ensures consistent citation resolution downstream, though page numbers
   won't match what you'd see if you opened the DOCX in Word.
"""
import io
import logging
from pathlib import Path
from typing import List, Tuple, TYPE_CHECKING

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph

from zipfile import ZipFile

try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    Paragraph = None  # type: ignore

from .models import ParsedDocument, ParsedPage, PageImage
from .ocr import ocr_image, is_tesseract_available

logger = logging.getLogger(__name__)


def extract_images_from_docx(file_path: Path) -> List[Tuple[str, bytes]]:
    """Extract embedded images from DOCX file using ZIP structure."""
    images = []
    
    try:
        with ZipFile(file_path, 'r') as docx_zip:
            # Images are stored in word/media/ directory
            for file_info in docx_zip.filelist:
                if file_info.filename.startswith('word/media/'):
                    image_bytes = docx_zip.read(file_info.filename)
                    images.append((file_info.filename, image_bytes))
    
    except Exception as e:
        logger.warning(f"Failed to extract images from DOCX: {e}")
    
    return images


def split_into_pseudo_pages(paragraphs: List["Paragraph"], max_chars_per_page: int = 3000) -> List[List["Paragraph"]]:
    """Split document paragraphs into pseudo-pages for provenance tracking."""
    if not paragraphs:
        return []
    
    pages = []
    current_page = []
    current_char_count = 0
    
    for para in paragraphs:
        text = para.text.strip()
        para_length = len(text)
        
        # Check if this is a heading (start new page)
        is_heading = para.style and 'Heading' in para.style.name
        
        # Start new page if:
        # 1. We have content and hit a heading
        # 2. We exceed max chars and have at least one paragraph
        if current_page and (is_heading or (current_char_count > max_chars_per_page)):
            pages.append(current_page)
            current_page = []
            current_char_count = 0
        
        # Add paragraph to current page
        current_page.append(para)
        current_char_count += para_length
    
    # Add final page
    if current_page:
        pages.append(current_page)
    
    # Ensure we have at least one page
    if not pages:
        pages = [[]]
    
    return pages


def parse_docx(file_path: str, filename: str, ocr_images: bool = True) -> ParsedDocument:
    """Parse a DOCX document and extract text and images."""
    if not PYTHON_DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not available. Please install: pip install python-docx"
        )
    
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
    
    ocr_available = ocr_images and is_tesseract_available()
    if ocr_images and not ocr_available:
        logger.warning("OCR for images requested but Tesseract is not available")
    
    try:
        # Open document
        doc = Document(file_path)
        
        # Extract all paragraphs (maintaining order)
        all_paragraphs = []
        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                all_paragraphs.append(para)
            elif isinstance(element, CT_Tbl):
                # Extract text from tables as well
                table = Table(element, doc)
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            all_paragraphs.append(para)
        
        # Split into pseudo-pages
        page_groups = split_into_pseudo_pages(all_paragraphs)
        
        # Extract images from DOCX ZIP structure
        docx_images = extract_images_from_docx(file_path)
        
        # Process each pseudo-page
        pages: List[ParsedPage] = []
        ocr_pages_used = 0
        
        for page_num, para_group in enumerate(page_groups):
            page_number = page_num + 1
            
            # Combine paragraph text
            text_parts = []
            for para in para_group:
                para_text = para.text.strip()
                if para_text:
                    # Preserve heading structure
                    if para.style and 'Heading' in para.style.name:
                        text_parts.append(f"\n{para_text}\n")
                    else:
                        text_parts.append(para_text)
            
            text = "\n".join(text_parts)
            
            # For this page, assign a subset of images
            # Simple strategy: divide images evenly across pages
            images_per_page = len(docx_images) // max(len(page_groups), 1)
            start_img_idx = page_num * images_per_page
            end_img_idx = start_img_idx + images_per_page if page_num < len(page_groups) - 1 else len(docx_images)
            
            page_images: List[PageImage] = []
            
            for img_idx in range(start_img_idx, end_img_idx):
                if img_idx < len(docx_images):
                    img_name, img_bytes = docx_images[img_idx]
                    
                    # Determine format from filename
                    img_format = img_name.split('.')[-1].upper() if '.' in img_name else 'UNKNOWN'
                    
                    # Optionally OCR the image
                    ocr_text = None
                    if ocr_available:
                        try:
                            ocr_text = ocr_image(img_bytes)
                            if ocr_text:
                                ocr_pages_used += 1
                        except Exception as e:
                            logger.debug(f"Image OCR failed for {img_name}: {e}")
                    
                    page_image = PageImage(
                        image_bytes=img_bytes,
                        page_number=page_number,
                        image_index=img_idx - start_img_idx,
                        format=img_format,
                        ocr_text=ocr_text
                    )
                    page_images.append(page_image)
            
            # Create parsed page
            parsed_page = ParsedPage(
                page_number=page_number,
                text=text,
                images=page_images,
                is_ocr=False,  # DOCX text is always native
                metadata={
                    "paragraph_count": len(para_group)
                }
            )
            pages.append(parsed_page)
        
        # Extract document core properties
        core_props = doc.core_properties
        metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or ""
        }
        
        return ParsedDocument(
            filename=filename,
            source_format="docx",
            pages=pages,
            total_pages=len(pages),
            ocr_pages_used=ocr_pages_used,
            metadata=metadata
        )
    
    except Exception as e:
        logger.error(f"Failed to parse DOCX {filename}: {e}")
        raise
